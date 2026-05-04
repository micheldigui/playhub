from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\PROJETOS\PLAYHUB")
OUT_DIR = ROOT / "docs" / "tdis"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "ETAPA_2_PLAYHUB_ABNT.docx"
FIG_PATH = OUT_DIR / "modelo_conceitual_playhub.png"


TITLE = (
    "Transformação Digital e Governança em Sistemas Esportivos Amadores: "
    "Avaliação de uma Plataforma Digital para Eficiência, Transparência e "
    "Experiência do Usuário"
)

DISCIPLINA = "Tecnologias Digitais na Indústria e Serviços"
CURSO = "Mestrado Profissional em Gestão e Tecnologia em Sistemas Produtivos"
AUTORES = "Manoel Denis Lopes do Nascimento\nMichel Silva de Souza"
PROFESSORES = (
    "Prof. Dr. Marcelo Duduchi Feitosa; Prof. Dra. Marcia Ito; "
    "Prof. Dr. Napoleão Galegale; Prof. Dr. Rafael Nobre Orsi; "
    "Prof. Dr. Waldemar Bonventi Jr."
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_cm):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C2D0"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_paragraph_font(paragraph, size=12, bold=None, italic=None, color=None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def format_body_paragraph(paragraph, first_line=True, justify=True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    if first_line:
        fmt.first_line_indent = Cm(1.25)
    else:
        fmt.first_line_indent = Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(paragraph, 12)


def add_body(doc, text, first_line=True):
    p = doc.add_paragraph()
    p.add_run(text)
    format_body_paragraph(p, first_line=first_line)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(p, 12, bold=True, color="1F2937")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = True
    set_paragraph_font(p, 10, bold=True)
    return p


def add_source(doc, text="Fonte: Elaborado pelos autores (2026)."):
    p = doc.add_paragraph()
    p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_font(p, 10)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    r_pr.append(r_fonts)
    r_pr.append(sz)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(r_pr)
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)
    set_paragraph_font(paragraph, 10)


def build_conceptual_model():
    width, height = 1800, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 42)
        font_box = ImageFont.truetype("arial.ttf", 30)
        font_small = ImageFont.truetype("arial.ttf", 24)
    except OSError:
        font_title = ImageFont.load_default()
        font_box = ImageFont.load_default()
        font_small = ImageFont.load_default()

    draw.text((width // 2, 45), "Modelo conceitual da pesquisa", anchor="mm", fill="#1F2937", font=font_title)

    def box(x1, y1, x2, y2, title, lines, fill, outline="#1F2937"):
        draw.rounded_rectangle([x1, y1, x2, y2], radius=18, fill=fill, outline=outline, width=4)
        draw.text(((x1 + x2) // 2, y1 + 36), title, anchor="mm", fill="#111827", font=font_box)
        y = y1 + 88
        for line in lines:
            draw.text(((x1 + x2) // 2, y), line, anchor="mm", fill="#111827", font=font_small)
            y += 36

    def arrow(x1, y1, x2, y2):
        draw.line([x1, y1, x2, y2], fill="#2563EB", width=8)
        if x2 > x1:
            pts = [(x2, y2), (x2 - 28, y2 - 16), (x2 - 28, y2 + 16)]
        else:
            pts = [(x2, y2), (x2 + 28, y2 - 16), (x2 + 28, y2 + 16)]
        draw.polygon(pts, fill="#2563EB")

    box(95, 260, 455, 510, "Variável independente", ["Adoção do PlayHub", "PWA, agenda, financeiro,", "governança e matchmaking"], "#E0F2FE")
    box(700, 155, 1115, 325, "Mecanismos", ["Automação de regras", "e coordenação digital"], "#EDE9FE")
    box(700, 380, 1115, 550, "Mecanismos", ["Transparência de dados", "e experiência de uso"], "#DCFCE7")
    box(1365, 120, 1710, 265, "VD1", ["Eficiência", "operacional"], "#FEEFAD")
    box(1365, 345, 1710, 490, "VD2", ["Transparência", "das informações"], "#FEEFAD")
    box(1365, 570, 1710, 715, "VD3", ["Experiência", "do usuário"], "#FEEFAD")
    box(660, 670, 1160, 790, "Variáveis de controle", ["perfil no grupo, tempo de participação, familiaridade digital"], "#F3F4F6")

    arrow(455, 385, 700, 240)
    arrow(455, 385, 700, 465)
    arrow(1115, 240, 1365, 190)
    arrow(1115, 240, 1365, 415)
    arrow(1115, 465, 1365, 415)
    arrow(1115, 465, 1365, 640)
    draw.line([910, 670, 910, 560], fill="#6B7280", width=5)
    draw.polygon([(910, 560), (895, 590), (925, 590)], fill="#6B7280")

    FIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    img.save(FIG_PATH)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E5E7EB")
        set_cell_margins(cell, 120, 110, 120, 110)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        set_paragraph_font(p, 9.5, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, 120, 110, 120, 110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.text = text
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            set_paragraph_font(p, 9.2)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        set_paragraph_font(p, 12)


def add_references(doc):
    refs = [
        "BUDLER, Marko; BOŽIČ, Katerina. Adopting transitional business models in small fitness businesses in response to business disruptions. Journal of Small Business Strategy, v. 33, n. 3, 2023. DOI: 10.53703/001c.92989.",
        "CHAI, Hongqin et al. Digital transformation in physical education: a bibliometric analysis of global trends, hotspots, and future directions through CiteSpace. SAGE Open, v. 15, 2025. DOI: 10.1177/21582440251395582.",
        "DE LEON, Frediezel et al. Digital transformation COVID-19 era: startup strategies for technology, management, and people. International Journal of Advances in Applied Sciences, v. 14, n. 1, p. 11-18, 2025. DOI: 10.11591/ijaas.v14.i1.pp11-18.",
        "ELSHAWESH, Yousef; WADA, Isah. Service quality and customer satisfaction: the mediating role of customer delight. International Review of Management and Marketing, v. 16, n. 2, p. 273-287, 2026. DOI: 10.32479/irmm.21607.",
        "HA, Taemin et al. A systematic review of technology-infused physical activity interventions in K-12 school settings: effectiveness, roles and implementation strategies. International Journal of Behavioral Nutrition and Physical Activity, v. 22, n. 113, 2025. DOI: 10.1186/s12966-025-01811-x.",
        "HASSAN, Ahmed K. et al. Enhancing operational performance: the impact of smart algorithms in Saudi Arabian sports facilities. Journal of Infrastructure, Policy and Development, v. 8, n. 8, 3628, 2024. DOI: 10.24294/jipd.v8i8.3628.",
        "MADSEN, Dag Øivind; GLEBOVA, Ekaterina. Sports Industry 5.0: reimagining sport through technology, humanity and sustainability. Frontiers in Sports and Active Living, v. 7, 1640362, 2025. DOI: 10.3389/fspor.2025.1640362.",
        "MAGAZ-GONZÁLEZ, Ana María et al. Technological structure and configuration of the use of technology in European sports organizations. Cultura, Ciencia y Deporte, v. 19, n. 60, 2145, 2024. DOI: 10.12800/ccd.v19i60.2145.",
        "MBANEFO, Chibuike C.; GROBBELAAR, Sara S. Unveiling the core elements of platform ecosystem development: a systemic lens for value co-creation in small and medium enterprises and orchestrators. Management Review Quarterly, v. 75, p. 1575-1618, 2025. DOI: 10.1007/s11301-024-00416-1.",
        "QUIÑONES, Daniela; ROJAS, Luis. Understanding the customer experience in human-computer interaction: a systematic literature review. PeerJ Computer Science, v. 9, e1219, 2023. DOI: 10.7717/peerj-cs.1219.",
        "STEGMANN, Pascal; STRÖBEL, Tim; WORATSCHEK, Herbert. Categorizing engagement behavior in sport brand communities: an empirical study informed by social practice theory. Sport Management Review, v. 27, n. 4, p. 544-571, 2024. DOI: 10.1080/14413523.2024.2329826.",
        "WANG, Chao; WANG, Zhigang. Influence of customer knowledge management on mobile fitness application customer value co-creation through flow experience and customer involvement. SAGE Open, v. 13, 2023. DOI: 10.1177/21582440231218777.",
        "WANG, Meng; WANG, Zheng; DENG, Rong. How to enhance Generation Z users' satisfaction experience with online fitness: a case study of fitness live streaming platforms. Frontiers in Computer Science, v. 6, 1499672, 2025. DOI: 10.3389/fcomp.2024.1499672.",
        "YANG, Jing; YAO, Tang; WANG, Jun. Diminishing returns of task-oriented interaction in digitally-mediated dynamic teams: evidence from amateur sports organizing. Frontiers in Psychology, v. 16, 1548846, 2025. DOI: 10.3389/fpsyg.2025.1548846.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.add_run(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        set_paragraph_font(p, 12)


def build_docx():
    build_conceptual_model()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    add_page_number(section.header.paragraphs[0])

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 1"].font.size = Pt(12)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True

    # Cover page.
    cover_items = [
        ("CENTRO ESTADUAL DE EDUCAÇÃO TECNOLÓGICA PAULA SOUZA", True, 12),
        ("UNIDADE DE PÓS-GRADUAÇÃO, EXTENSÃO E PESQUISA", True, 12),
        (CURSO.upper(), True, 12),
        ("", False, 12),
        ("RELATÓRIO DE ACOMPANHAMENTO DO ARTIGO - ETAPA 2", True, 12),
        (DISCIPLINA.upper(), True, 12),
    ]
    for text, bold, size in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = bold
        set_paragraph_font(p, size)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(TITLE.upper())
    r.bold = True
    set_paragraph_font(p, 12)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.add_run("Discentes:\n").bold = True
    p.add_run(AUTORES)
    set_paragraph_font(p, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.add_run("Professores:\n").bold = True
    p.add_run(PROFESSORES)
    set_paragraph_font(p, 12)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("São Paulo\n2026")
    set_paragraph_font(p, 12)
    doc.add_page_break()

    add_heading(doc, "1 CONTINUIDADE DA ETAPA 1 E DELIMITAÇÃO DO ESTUDO")
    add_body(
        doc,
        "A Etapa 1 definiu como problema de pesquisa a análise do impacto da adoção de uma plataforma digital baseada em governança e automação, o PlayHub, sobre a eficiência operacional, a transparência das informações e a experiência dos usuários em grupos esportivos amadores. Esta Etapa 2 dá continuidade a essa formulação, mantendo a questão central e refinando a hipótese geral em constructos mensuráveis, referencial teórico e metodologia de pesquisa.",
    )
    add_body(
        doc,
        "O objeto empírico permanece sendo um sistema produtivo informal de serviços esportivos. Nesse contexto, a prática esportiva depende de atividades recorrentes de coordenação, confirmação de presença, cobrança, controle de regras, comunicação e recomposição de vagas. Antes do PlayHub, tais rotinas tendiam a ocorrer por aplicativos genéricos de mensagem e controles manuais, criando dispersão informacional, baixa rastreabilidade e sobrecarga para gestores voluntários.",
    )
    add_body(
        doc,
        "O ajuste conceitual proposto nesta etapa é tratar o PlayHub não apenas como aplicativo operacional, mas como uma intervenção de transformação digital em micro-sistemas de serviço. A literatura recente compreende a transformação digital como a integração entre tecnologia, gestão/processos e pessoas para redefinir proposições de valor e melhorar operações (DE LEON et al., 2025). Essa leitura é aderente ao caso, pois a plataforma altera simultaneamente a ferramenta utilizada, a forma de governança do grupo e o papel dos participantes na organização do serviço.",
    )
    add_body(
        doc,
        "Também se propõe substituir expressões excessivamente promocionais, como “sistema revolucionário” ou “punição algorítmica”, por termos academicamente mais precisos: governança digital assistida, automação de regras, coordenação digital e ecossistema de plataforma. Essa mudança preserva a ideia original da Etapa 1, mas melhora a aderência ao padrão científico esperado em um trabalho de mestrado.",
    )

    add_heading(doc, "2 OBJETIVOS")
    add_heading(doc, "2.1 Objetivo geral", level=2)
    add_body(
        doc,
        "Analisar o impacto da adoção da plataforma digital PlayHub na eficiência operacional, na transparência das informações e na experiência dos usuários em grupos esportivos amadores.",
    )
    add_heading(doc, "2.2 Objetivos específicos", level=2)
    add_bullets(
        doc,
        [
            "Caracterizar o processo de organização de grupos esportivos amadores antes da adoção do PlayHub, identificando atividades manuais, pontos de atrito e limitações de transparência.",
            "Mapear as funcionalidades do PlayHub relacionadas à governança digital, à gestão financeira, à agenda de partidas, à confirmação de presença, ao perfil esportivo e ao matchmaking entre atletas e equipes.",
            "Mensurar a variação percebida na eficiência operacional dos administradores após a adoção da plataforma, considerando tempo de gestão, esforço de coordenação e carga cognitiva.",
            "Avaliar a percepção dos usuários quanto à transparência das informações, especialmente regras, status financeiro, critérios de presença e comunicação das decisões do grupo.",
            "Analisar a experiência dos usuários com a plataforma, considerando facilidade de uso, clareza da interface, satisfação e intenção de recomendação.",
            "Relacionar os resultados empíricos aos constructos identificados na literatura sobre transformação digital, ecossistemas de plataforma, coordenação de equipes mediada digitalmente e experiência do usuário.",
        ],
    )

    add_heading(doc, "3 BIBLIOMETRIA E SELEÇÃO DO REFERENCIAL TEÓRICO")
    add_body(
        doc,
        "A bibliometria desta etapa foi organizada como um levantamento inicial do acervo bibliográfico disponível para o projeto. O objetivo não é substituir uma revisão sistemática completa, mas demonstrar rastreabilidade na seleção dos artigos que fundamentam o modelo conceitual e a metodologia. O corpus local contém 94 arquivos científicos em PDF, sendo 26 classificados como referencial direto e 68 como apoio complementar. Para a Etapa 2, foram priorizados artigos recentes, publicados entre 2023 e 2026, com aderência direta ao problema de pesquisa.",
    )

    add_caption(doc, "Quadro 1 - Síntese bibliométrica do acervo analisado")
    add_table(
        doc,
        ["Critério", "Resultado adotado na Etapa 2"],
        [
            ["Corpus local identificado", "94 publicações em PDF, separadas em referencial direto e referencial complementar."],
            ["Recorte principal", "26 publicações classificadas como Referencial Direto (Core)."],
            ["Período dominante", "2023 a 2026, com ênfase em literatura recente sobre transformação digital, esporte e serviços digitais."],
            ["Descritores temáticos", "digital transformation; sport; sports organizations; platform ecosystem; value co-creation; human-computer interaction; customer experience; smart algorithms; amateur sports organizing; mobile fitness applications."],
            ["Critérios de inclusão", "Aderência ao PlayHub, publicação em periódico científico, relação com variáveis do modelo e contribuição para discutir eficiência, transparência ou experiência do usuário."],
            ["Critérios de exclusão", "Materiais com baixa relação com o problema, duplicações temáticas e documentos sem contribuição direta para o constructo da Etapa 2."],
        ],
        [4.5, 10.5],
    )
    add_source(doc)

    add_caption(doc, "Quadro 2 - Artigos centrais selecionados e contribuição para o PlayHub")
    add_table(
        doc,
        ["Eixo", "Autores", "Contribuição para o trabalho"],
        [
            ["Transformação digital em serviços", "De Leon et al. (2025)", "Define transformação digital como articulação entre tecnologia, gestão/processos e pessoas, eixo que organiza a leitura do PlayHub."],
            ["Pequenos negócios e resiliência", "Budler e Božič (2023)", "Sustenta a transição de modelos operacionais tradicionais para modelos digitais ajustáveis em pequenos negócios fitness."],
            ["Organizações esportivas e digitalização", "Magaz-González et al. (2024)", "Mostra que organizações esportivas ainda enfrentam baixo uso tecnológico e necessidade de cultura digital."],
            ["Ecossistema de plataforma", "Mbanefo e Grobbelaar (2025)", "Fundamenta o PlayHub como ecossistema de múltiplos atores com co-criação de valor e necessidade de governança transparente."],
            ["Prática social em comunidades esportivas", "Stegmann, Ströbel e Woratschek (2024)", "Permite interpretar engajamento, pertencimento e rotinas digitais como práticas sociais em comunidades esportivas."],
            ["Coordenação em equipes digitais", "Yang, Yao e Wang (2025)", "Aproxima diretamente o estudo de plataformas para organização de esportes amadores e eficiência de times dinâmicos."],
            ["Algoritmos e desempenho operacional", "Hassan et al. (2024)", "Apoia a análise de automação de regras e sua relação com desempenho operacional em instalações esportivas."],
            ["Experiência do usuário em HCI", "Quiñones e Rojas (2023)", "Fornece base para avaliar experiência do usuário em produtos de software como fenômeno multidimensional."],
            ["Aplicativos fitness e co-criação", "Wang e Wang (2023)", "Relaciona gestão de conhecimento, envolvimento do usuário, experiência de fluxo e co-criação de valor em apps fitness."],
            ["Satisfação em plataformas fitness", "Wang, Wang e Deng (2025)", "Indica que interação social, utilidade, conveniência e qualidade técnica influenciam satisfação em plataformas digitais fitness."],
            ["Tecnologia e atividade física", "Ha et al. (2025)", "Mostra que aplicativos e plataformas web atuam como ferramentas de comunicação e intervenção em atividades físicas."],
            ["Esporte 5.0", "Madsen e Glebova (2025)", "Conecta tecnologia, centralidade humana, inclusão e resiliência no esporte, fortalecendo a contribuição social do PlayHub."],
        ],
        [3.7, 4.2, 7.1],
    )
    add_source(doc)

    add_heading(doc, "4 REFERENCIAL TEÓRICO")
    add_heading(doc, "4.1 Transformação digital em micro-sistemas de serviços", level=2)
    add_body(
        doc,
        "A transformação digital não deve ser reduzida à adoção de uma ferramenta tecnológica isolada. A literatura recente a descreve como um processo de mudança que envolve tecnologias digitais, gestão, processos e pessoas, com potencial de melhorar produtividade, reduzir custos e redefinir a entrega de valor (DE LEON et al., 2025). Essa abordagem permite compreender o PlayHub como uma intervenção sistêmica, pois a plataforma reorganiza fluxos que antes dependiam de mensagens dispersas, memória individual e decisões informais.",
    )
    add_body(
        doc,
        "Em pequenos negócios fitness, a adoção de modelos transicionais revelou-se uma resposta a situações de instabilidade, permitindo ajustes temporários ou graduais na proposição de valor e nos processos de serviço (BUDLER; BOŽIČ, 2023). Embora o PlayHub esteja situado em grupos esportivos amadores, há uma analogia relevante: ambos operam com recursos limitados, forte dependência de relacionamento com usuários e necessidade de continuidade operacional. A digitalização, nesse caso, passa a funcionar como mecanismo de resiliência e padronização do serviço.",
    )
    add_body(
        doc,
        "O estudo de organizações esportivas europeias evidencia que o setor esportivo ainda apresenta desafios de digitalização, especialmente em competências digitais, uso de dados e estrutura tecnológica (MAGAZ-GONZÁLEZ et al., 2024). Essa constatação reforça a relevância acadêmica do PlayHub: a pesquisa observa uma transformação digital em escala micro, em um ambiente ainda pouco estruturado, mas socialmente recorrente.",
    )

    add_heading(doc, "4.2 Plataformas digitais, ecossistemas e co-criação de valor", level=2)
    add_body(
        doc,
        "Plataformas digitais são relevantes quando deixam de atuar apenas como ferramentas internas e passam a coordenar diferentes atores em um ecossistema. Mbanefo e Grobbelaar (2025) destacam que a conceituação de valor se deslocou de uma perspectiva centrada na firma para uma perspectiva ecossistêmica, envolvendo atores que integram recursos e co-criam valor. No PlayHub, esse ecossistema é composto por administradores, atletas fixos, atletas avulsos, equipes públicas ou privadas e rotinas de governança.",
    )
    add_body(
        doc,
        "A co-criação ocorre quando os usuários não apenas consomem a plataforma, mas alimentam o sistema com dados, decisões e interações: confirmação de presença, pagamentos, solicitações de ingresso, convites, perfil esportivo, preferências de modalidade e sinais de disponibilidade. Essa dinâmica é coerente com estudos sobre aplicativos fitness, nos quais conhecimento do cliente, envolvimento e experiência de fluxo influenciam a co-criação de valor (WANG; WANG, 2023).",
    )
    add_body(
        doc,
        "Nas comunidades esportivas, a interação digital também reforça práticas de engajamento. Stegmann, Ströbel e Woratschek (2024), a partir da Teoria da Prática Social, demonstram que plataformas digitais podem estruturar práticas de participação, pertencimento e sustentação de comunidades esportivas. No PlayHub, a prática social do grupo passa a ser mediada por regras, perfis, listas e informações compartilhadas, o que reduz ambiguidade e fortalece previsibilidade.",
    )

    add_heading(doc, "4.3 Coordenação digital em equipes esportivas amadoras", level=2)
    add_body(
        doc,
        "A pesquisa de Yang, Yao e Wang (2025) é central para este trabalho por tratar de equipes dinâmicas mediadas digitalmente no contexto de organização de esportes amadores. Os autores mostram que a eficiência de organização de equipes depende da qualidade das interações orientadas à tarefa e da identificação dos membros com o time. Essa discussão se aproxima diretamente do PlayHub, pois a plataforma busca reduzir o excesso de mensagens operacionais e converter a coordenação em fluxos mais claros e rastreáveis.",
    )
    add_body(
        doc,
        "Essa base permite argumentar que a digitalização não deve simplesmente aumentar a quantidade de comunicação. O ganho esperado está em transformar comunicação dispersa em informação estruturada. Assim, funcionalidades como agenda, lista de presença, fila de espera, status financeiro e convites digitais devem ser avaliadas pelo quanto reduzem ruído, retrabalho e decisões informais, e não apenas pelo número de interações geradas.",
    )

    add_heading(doc, "4.4 Governança digital assistida e automação de regras", level=2)
    add_body(
        doc,
        "A governança digital assistida é uma das dimensões centrais do PlayHub. No lugar de regras aplicadas de forma pessoal pelo gestor, o sistema torna visíveis critérios e consequências, como status financeiro, confirmação de presença, cartões, strikes e restrições temporárias. Hassan et al. (2024) apontam correlação entre uso de algoritmos inteligentes e desempenho operacional em instalações esportivas, sugerindo que automações podem apoiar decisão, planejamento, análise de dados e solução de problemas.",
    )
    add_body(
        doc,
        "No caso estudado, o objetivo não é delegar julgamento moral ao software, mas automatizar regras previamente pactuadas pelo grupo. Por isso, o termo mais adequado é governança digital assistida: o algoritmo executa critérios definidos, enquanto os administradores permanecem responsáveis pela configuração das regras, pelo acompanhamento de exceções e pela mediação quando necessário. Essa distinção evita uma leitura tecnicista e preserva a centralidade humana da gestão.",
    )
    add_body(
        doc,
        "A discussão também dialoga com a proposta de Sports Industry 5.0, que defende tecnologia orientada por centralidade humana, inclusão, sustentabilidade e resiliência no esporte (MADSEN; GLEBOVA, 2025). O PlayHub se aproxima desse paradigma ao buscar eficiência sem eliminar o papel comunitário do esporte amador.",
    )

    add_heading(doc, "4.5 Experiência do usuário, HCI e satisfação", level=2)
    add_body(
        doc,
        "A experiência do usuário é um constructo multidimensional. Quiñones e Rojas (2023), em revisão sistemática sobre customer experience e interação humano-computador, mostram que a experiência em produtos de software envolve dimensões cognitivas, emocionais, funcionais e contextuais. Para o PlayHub, isso significa que a avaliação não deve considerar apenas se a funcionalidade existe, mas se o usuário compreende o fluxo, confia na informação e percebe redução de esforço.",
    )
    add_body(
        doc,
        "Em plataformas fitness digitais, a satisfação é influenciada por interação social, utilidade, conveniência, qualidade funcional e qualidade técnica (WANG; WANG; DENG, 2025). Esses elementos são compatíveis com os módulos do PlayHub: o app precisa ser útil para organizar o jogo, conveniente por funcionar como PWA, tecnicamente estável e socialmente capaz de aproximar atletas e equipes. Estudos sobre intervenções de atividade física com tecnologia também indicam que aplicativos e plataformas web podem atuar como componentes de comunicação e intervenção, desde que bem desenhados para o contexto de uso (HA et al., 2025).",
    )
    add_body(
        doc,
        "Assim, a experiência do usuário será tratada nesta pesquisa como percepção de facilidade, clareza, satisfação e recomendação, combinando indicadores de uso com avaliação subjetiva dos participantes.",
    )

    add_heading(doc, "5 MODELO CONCEITUAL E HIPÓTESES")
    add_body(
        doc,
        "A hipótese geral da Etapa 1 permanece válida: a adoção da plataforma PlayHub melhora a eficiência operacional, a transparência das informações e a experiência dos usuários em comparação aos métodos tradicionais. Para a Etapa 2, essa hipótese foi desdobrada em hipóteses operacionais, permitindo melhor mensuração na etapa de coleta e análise de dados.",
    )

    add_caption(doc, "Figura 1 - Modelo conceitual da pesquisa")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(FIG_PATH), width=Cm(15.0))
    add_source(doc)

    add_heading(doc, "5.1 Hipóteses", level=2)
    add_bullets(
        doc,
        [
            "H1: a adoção do PlayHub reduz o tempo e o esforço percebido de gestão operacional dos administradores de grupos esportivos amadores.",
            "H2: a adoção do PlayHub aumenta a percepção de transparência das informações organizacionais, especialmente quanto a regras, presenças e finanças.",
            "H3: a adoção do PlayHub melhora a experiência dos usuários, medida por facilidade de uso, clareza, satisfação e intenção de recomendação.",
            "H4: a governança digital assistida e a coordenação digital atuam como mecanismos explicativos entre a adoção da plataforma e os resultados percebidos.",
        ],
    )

    add_caption(doc, "Quadro 3 - Relação entre constructos, variáveis e métricas")
    add_table(
        doc,
        ["Constructo", "Tipo", "Indicadores propostos", "Escala/Fonte"],
        [
            ["Adoção do PlayHub", "Variável independente", "Uso dos módulos de agenda, presença, financeiro, perfil esportivo, convites, notificações e governança.", "Registro de uso e questionário pós-adoção."],
            ["Eficiência operacional", "Variável dependente 1", "Tempo semanal de gestão; quantidade de tarefas manuais; esforço percebido; redução de retrabalho.", "Minutos/semana e escala Likert de 1 a 5."],
            ["Transparência das informações", "Variável dependente 2", "Clareza das regras; visibilidade financeira; previsibilidade de critérios; confiança na informação.", "Escala Likert de 1 a 5."],
            ["Experiência do usuário", "Variável dependente 3", "Facilidade de uso; clareza da interface; satisfação; intenção de recomendação.", "Escala Likert de 1 a 5 e NPS de 0 a 10."],
            ["Perfil do participante", "Controle", "Administrador ou atleta; tempo no grupo; familiaridade digital; modalidade praticada.", "Questões de caracterização."],
        ],
        [3.4, 3.0, 5.5, 3.1],
    )
    add_source(doc)

    add_heading(doc, "6 METODOLOGIA")
    add_heading(doc, "6.1 Classificação da pesquisa", level=2)
    add_body(
        doc,
        "A pesquisa é classificada como científica aplicada, pois busca produzir conhecimento sobre o impacto de tecnologias digitais em um contexto real de serviços esportivos amadores, com potencial de orientar decisões práticas de gestão e desenvolvimento de plataformas. Embora exista desenvolvimento tecnológico associado ao PlayHub, o foco do artigo não é apenas relatar a construção do software, mas testar hipóteses sobre seus efeitos percebidos.",
    )
    add_body(
        doc,
        "Quanto ao objetivo, a pesquisa é explicativa, com componentes descritivos. É explicativa porque busca analisar a relação entre adoção da plataforma e variações em eficiência, transparência e experiência do usuário. É descritiva porque também caracteriza o contexto, o perfil dos participantes e as práticas de organização antes e depois da adoção.",
    )
    add_body(
        doc,
        "A abordagem é mista, com predominância quantitativa. Os dados principais serão obtidos por questionários estruturados em escala Likert e indicadores de tempo, permitindo comparação entre médias antes e depois da adoção. Comentários abertos e observações dos administradores serão utilizados de forma complementar para interpretar os resultados.",
    )

    add_heading(doc, "6.2 Delineamento", level=2)
    add_body(
        doc,
        "O delineamento adotado é um estudo de caso único com desenho quase-experimental pré e pós-implementação, sem grupo controle. A escolha se justifica porque o PlayHub será analisado em um ambiente real, no qual não é viável controlar todas as variáveis externas, mas é possível comparar percepções e indicadores antes e depois da intervenção.",
    )
    add_body(
        doc,
        "O procedimento técnico combina pesquisa bibliográfica, levantamento por questionário e análise documental/operacional da plataforma. A pesquisa bibliográfica sustenta o modelo conceitual; o levantamento coleta percepções dos usuários; e a análise documental/operacional considera fluxos e registros do sistema, como presença, pagamentos, convites, notificações e configuração de regras.",
    )

    add_heading(doc, "6.3 Universo, amostra e participantes", level=2)
    add_body(
        doc,
        "O universo da pesquisa é composto por grupos esportivos amadores que dependem de coordenação recorrente para realizar partidas, administrar contribuições financeiras e organizar participantes. A unidade de análise inicial é o grupo piloto usuário do PlayHub.",
    )
    add_body(
        doc,
        "A amostra é não probabilística, por conveniência e intencionalidade, composta por administradores e atletas que utilizarem a plataforma durante o período de observação. A rodada piloto prevista considera participantes com papéis distintos - administradores e jogadores - para permitir comparação entre a percepção de quem gerencia e a de quem consome o serviço. Caso novos grupos sejam incorporados até a Etapa 3, a amostra poderá ser ampliada mantendo os mesmos instrumentos.",
    )

    add_heading(doc, "6.4 Procedimentos de coleta de dados", level=2)
    add_body(
        doc,
        "A coleta será realizada em três momentos. No primeiro, será aplicado um questionário diagnóstico sobre o modelo anterior de organização, baseado em WhatsApp, planilhas ou controles manuais. No segundo, os usuários utilizarão o PlayHub em um ciclo operacional mínimo, envolvendo confirmação de presença, consulta de informações, gestão financeira e comunicação de regras. No terceiro, será aplicado questionário pós-adoção para medir as mesmas dimensões e captar a experiência de uso.",
    )
    add_body(
        doc,
        "Os questionários utilizarão escala Likert de cinco pontos para medir clareza, facilidade, transparência e esforço percebido. Para recomendação, será utilizada escala de 0 a 10. Administradores também informarão tempo aproximado gasto em rotinas de gestão antes e depois da adoção. Sempre que possível, esses dados serão comparados com evidências operacionais do sistema, como registros de presença, ciclo financeiro e uso de notificações.",
    )

    add_caption(doc, "Quadro 4 - Instrumento de coleta proposto")
    add_table(
        doc,
        ["Dimensão", "Exemplos de itens do questionário", "Respondentes"],
        [
            ["Eficiência operacional", "Quanto tempo semanal é gasto com listas, cobranças e resolução de dúvidas? O processo de gestão exige esforço excessivo?", "Administradores."],
            ["Transparência", "As regras do grupo são claras? O status financeiro é visível? Os critérios de presença e fila são compreensíveis?", "Administradores e atletas."],
            ["Experiência do usuário", "O sistema é fácil de usar? As informações são fáceis de encontrar? Você recomendaria o PlayHub a outro grupo?", "Administradores e atletas."],
            ["Governança digital", "As regras automatizadas reduzem conflitos? A aplicação de critérios pelo sistema é percebida como consistente?", "Administradores e atletas."],
            ["Caracterização", "Papel no grupo, tempo de participação, modalidade, idade aproximada e familiaridade com ferramentas digitais.", "Administradores e atletas."],
        ],
        [3.3, 8.0, 3.7],
    )
    add_source(doc)

    add_heading(doc, "6.5 Tratamento e análise dos dados", level=2)
    add_body(
        doc,
        "Os dados serão analisados por estatística descritiva, incluindo média, mediana, frequência, variação percentual e comparação de médias pré e pós-adoção. Para amostras pequenas, a interpretação priorizará a magnitude da diferença e a coerência com o referencial teórico. Caso a amostra seja ampliada, poderão ser aplicados testes não paramétricos, como Wilcoxon para dados pareados, ou teste t pareado quando os pressupostos forem atendidos.",
    )
    add_body(
        doc,
        "A análise qualitativa complementar será aplicada a respostas abertas e observações dos administradores, buscando explicar por que determinados indicadores melhoraram, permaneceram estáveis ou pioraram. Essa etapa será utilizada para discutir resultados à luz de coordenação digital, governança assistida e experiência do usuário, conforme os autores do referencial teórico.",
    )

    add_heading(doc, "6.6 Confiabilidade, validade e limitações", level=2)
    add_body(
        doc,
        "A confiabilidade será buscada por meio da padronização dos instrumentos de coleta, aplicação das mesmas dimensões no pré e no pós, registro dos critérios de análise e manutenção de evidências operacionais da plataforma. A validade interna será limitada pela ausência de grupo controle, mas fortalecida pela comparação temporal e pela triangulação entre percepção dos usuários, indicadores de tempo e registros do sistema.",
    )
    add_body(
        doc,
        "A validade externa é limitada por se tratar de estudo de caso único em um grupo piloto. Ainda assim, a pesquisa pode gerar contribuição analítica para contextos similares de serviços esportivos amadores, especialmente aqueles com baixa estrutura formal, dependência de gestores voluntários e uso de mensageria como principal ferramenta de coordenação.",
    )

    add_heading(doc, "7 CONTRIBUIÇÃO ESPERADA")
    add_body(
        doc,
        "A contribuição acadêmica esperada é demonstrar como a transformação digital pode ser analisada em sistemas produtivos informais e de pequena escala, ampliando a discussão para além de indústrias, grandes organizações e serviços altamente formalizados. Ao articular transformação digital, ecossistemas de plataforma, coordenação de equipes esportivas e experiência do usuário, o estudo propõe um modelo aplicável ao esporte amador.",
    )
    add_body(
        doc,
        "A contribuição prática é oferecer evidências para o desenvolvimento de soluções digitais que reduzam sobrecarga operacional, aumentem transparência e melhorem a experiência de usuários em comunidades esportivas. Para gestores voluntários, a plataforma pode reduzir atividades repetitivas e conflitos associados à aplicação manual de regras. Para atletas, pode melhorar acesso à informação, previsibilidade e participação.",
    )

    add_heading(doc, "8 CONSIDERAÇÕES PARCIAIS DA ETAPA 2")
    add_body(
        doc,
        "A Etapa 2 consolidou a continuidade da pesquisa iniciada na Etapa 1, mantendo a questão central e a hipótese geral, mas refinando a estrutura científica do trabalho. O PlayHub passa a ser analisado como intervenção de transformação digital em um sistema produtivo informal de serviços esportivos, com variáveis dependentes claramente definidas: eficiência operacional, transparência das informações e experiência do usuário.",
    )
    add_body(
        doc,
        "O referencial teórico selecionado fortalece a pesquisa por conectar o caso a discussões atuais sobre plataformas digitais, co-criação de valor, práticas sociais em comunidades esportivas, coordenação digital de equipes amadoras, governança por algoritmos e experiência do usuário em sistemas computacionais. A próxima etapa deverá executar os procedimentos de coleta, organizar os dados em tabelas e gráficos e iniciar a análise dos resultados à luz das hipóteses aqui formuladas.",
    )

    doc.add_page_break()
    add_heading(doc, "REFERÊNCIAS")
    add_references(doc)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Relatório de Acompanhamento do Artigo - Etapa 2"
    doc.core_properties.author = "Manoel Denis Lopes do Nascimento; Michel Silva de Souza"
    doc.core_properties.keywords = "PlayHub; transformação digital; governança digital; esporte amador; ABNT"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
